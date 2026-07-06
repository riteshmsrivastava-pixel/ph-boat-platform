#!/usr/bin/env python3
"""Generate an editable Google-Docs-ready .docx mirroring the PH Boat pack.

Parses the in-repo design source so the doc matches the live site (prose + tables,
inline bold/italic, lists). Figures (SVG) become a placeholder pointing at the site.
Team edits this doc in Google Docs; the site stays the presentation layer and is
re-synced on request.
"""
import re, os, html
from html.parser import HTMLParser
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = os.path.expanduser("~/ph-boat-platform/design-source.dc.html")
OUT = os.path.expanduser("~/Desktop/PH-Boat-Pack-EDITABLE.docx")
LIVE = "https://riteshmsrivastava-pixel.github.io/ph-boat-platform/"

raw = open(SRC, encoding="utf-8").read()
raw = re.sub(r'\s*style-hover="[^"]*"', "", raw)
main_inner = re.search(r"<main\b[^>]*>(.*)</main>", raw, re.S).group(1)

# ---------- tiny DOM ----------
class N:
    def __init__(self, tag, attrs=None):
        self.tag = tag; self.attrs = dict(attrs or []); self.children = []; self.text = None
    @property
    def is_text(self): return self.tag is None

VOID = {"br", "img", "hr", "meta", "link", "input"}
INLINE_KEEP = {"strong", "b", "em", "i", "a", "span", "br"}

class DOM(HTMLParser):
    def __init__(self):
        super().__init__(convert_charrefs=True)
        self.root = N("root"); self.stack = [self.root]
    def handle_starttag(self, tag, attrs):
        node = N(tag, attrs); self.stack[-1].children.append(node)
        if tag not in VOID: self.stack.append(node)
    def handle_startendtag(self, tag, attrs):
        self.stack[-1].children.append(N(tag, attrs))
    def handle_endtag(self, tag):
        for i in range(len(self.stack) - 1, 0, -1):
            if self.stack[i].tag == tag:
                del self.stack[i:]; break
    def handle_data(self, data):
        t = re.sub(r"\s+", " ", data)
        if t:
            n = N(None); n.text = t; self.stack[-1].children.append(n)

dom = DOM(); dom.feed(main_inner)

# strip em-dashes everywhere (standing preference)
def _desh_walk(node):
    if node.is_text and node.text:
        node.text = re.sub(r"[ \t]*—[ \t]*", " - ", node.text)
    for c in node.children:
        _desh_walk(c)
_desh_walk(dom.root)

BLOCK = {"h1", "h2", "h3", "h4", "p", "table", "ul", "ol", "dl", "div", "figure"}

def text_of(node):
    if node.is_text: return node.text or ""
    return "".join(text_of(c) for c in node.children)

# ---------- group top-level section/figure nodes into pages ----------
top = [c for c in dom.root.children if c.tag in ("section", "figure")]
PAGES = [
    ("s1", "01", "Problem"),         ("s2", "02", "Research & Data"),
    ("s3", "03", "Personas"),        ("s4", "04", "Interview Guide"),
    ("s5", "05", "Econ Concepts"),   ("s6", "06", "Business Model"),
    ("s7", "07", "Slides"),
]
pages, cur = {}, None
for node in top:
    sid = node.attrs.get("id")
    if sid in dict((p[0], 1) for p in PAGES):
        cur = sid; pages[sid] = [node]
    elif cur:
        pages[cur].append(node)

# ---------- docx build ----------
doc = Document()
doc.styles["Normal"].font.name = "Georgia"
doc.styles["Normal"].font.size = Pt(11)

def shade(cell, color, white=False):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd"); sh.set(qn("w:val"), "clear"); sh.set(qn("w:fill"), color)
    tcPr.append(sh)

def add_inline(p, node, bold=False, italic=False):
    for c in node.children:
        if c.is_text:
            if c.text:
                r = p.add_run(c.text)
                if bold: r.bold = True
                if italic: r.italic = True
        elif c.tag == "br":
            p.add_run().add_break()
        elif c.tag == "span" and not text_of(c).strip():
            continue  # decorative rule span
        else:
            add_inline(p, c,
                       bold or c.tag in ("strong", "b"),
                       italic or c.tag in ("em", "i"))

def render_table(tnode):
    rows = [r for r in tnode.children if r.tag == "tr"]
    if not rows: return
    ncols = 0
    for r in rows:
        cells = [c for c in r.children if c.tag in ("td", "th")]
        ncols = max(ncols, sum(int(c.attrs.get("colspan", 1)) for c in cells))
    t = doc.add_table(rows=0, cols=ncols); t.style = "Table Grid"
    for r in rows:
        cells = [c for c in r.children if c.tag in ("td", "th")]
        is_head = any(c.tag == "th" for c in cells)
        row = t.add_row()
        ci = 0
        for c in cells:
            span = int(c.attrs.get("colspan", 1))
            cell = row.cells[ci]
            if span > 1:
                cell = cell.merge(row.cells[ci + span - 1])
            cell.paragraphs[0].text = ""
            add_inline(cell.paragraphs[0], c, bold=is_head)
            if is_head:
                shade(cell, "162A34")
                for rn in cell.paragraphs[0].runs:
                    rn.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF); rn.font.size = Pt(9)
            elif "background:#0C1C26" in c.attrs.get("style", ""):
                shade(cell, "0C1C26")
                for rn in cell.paragraphs[0].runs:
                    rn.font.color.rgb = RGBColor(0xEF, 0xE6, 0xD3)
            elif "background:#EFE7D7" in c.attrs.get("style", ""):
                shade(cell, "EFE7D7")
            ci += span
    doc.add_paragraph()

def render_block(node):
    tag = node.tag
    if tag == "header":
        return  # section number/title handled by caller
    if tag == "h3":
        doc.add_heading(text_of(node).strip(), level=2)
    elif tag == "h4":
        doc.add_heading(text_of(node).strip(), level=3)
    elif tag == "p":
        p = doc.add_paragraph(); add_inline(p, node)
    elif tag == "table":
        render_table(node)
    elif tag in ("ul", "ol"):
        style = "List Number" if tag == "ol" else "List Bullet"
        for li in [c for c in node.children if c.tag == "li"]:
            p = doc.add_paragraph(style=style); add_inline(p, li)
    elif tag == "dl":
        cur = None
        for c in node.children:
            if c.tag == "dt":
                cur = doc.add_paragraph(); cur.add_run(text_of(c).strip() + ": ").bold = True
            elif c.tag == "dd":
                if cur is None: cur = doc.add_paragraph()
                add_inline(cur, c); cur = None
    elif tag == "div":
        if any((not c.is_text) and c.tag in BLOCK for c in node.children):
            for c in node.children:      # container -> render children as blocks
                if not c.is_text: render_block(c)
        else:                            # leaf callout / label -> one paragraph
            p = doc.add_paragraph(); add_inline(p, node)
            p.paragraph_format.left_indent = Pt(14)
    elif tag == "figure":
        cap = ""
        for d in node.children:
            m = re.search(r"Figure 0(\d)", text_of(d))
            if m: cap = "Figure " + m.group(1); break
        doc.add_paragraph(
            f"[ {cap or 'Figure'}: chart lives on the live site - not editable here. See {LIVE} ]"
        ).runs[0].italic = True

# ---- intro / how-to ----
h = doc.add_heading("Verified Small-Boat Transfers in the Philippines", level=0)
doc.add_paragraph("15.024 Applied Economics for Managers - Team Project Pack - EDITABLE SOURCE").runs[0].italic = True

doc.add_heading("How to use this document", level=1)
for lead, rest in [
    ("", "This is our shared, editable source of truth for the project pack. Edit freely, leave comments, or use Suggesting mode if you want changes reviewed first."),
    ("Live site (read / present view): ", LIVE),
    ("Structure: ", "each Section below maps one-to-one to a page on the site."),
    ("Publishing: ", "when we want the site refreshed to match this doc, ping Gaurav - it re-publishes in about a minute."),
    ("Charts: ", "the two figures (price ladder, model diagram) live on the site only and are marked with a [Figure] note here; everything else is editable in this doc."),
    ("Rigor: ", "keep numbers tagged SOURCED vs ESTIMATED honest, and flag anything that should be re-verified before it hits a slide."),
]:
    p = doc.add_paragraph(style="List Bullet")
    if lead:
        p.add_run(lead).bold = True
    p.add_run(rest)

doc.add_paragraph()
doc.add_paragraph("--- The pack begins on the next page ---").runs[0].italic = True
doc.add_page_break()

# ---- sections ----
for sid, num, name in PAGES:
    first = True
    for node in pages[sid]:
        if node.tag == "section":
            header = next((c for c in node.children if c.tag == "header"), None)
            if header:
                h2txt = ""
                def find_h2(n):
                    global h2txt
                    for c in n.children:
                        if c.tag == "h2": h2txt = text_of(c)
                        elif not c.is_text: find_h2(c)
                find_h2(header)
                doc.add_heading(f"Section {num} - {h2txt.strip()}", level=1)
            for c in node.children:
                if c.tag != "header":
                    render_block(c)
        elif node.tag == "figure":
            render_block(node)
    if sid != "s7":
        doc.add_page_break()

doc.save(OUT)
print("saved", OUT)
